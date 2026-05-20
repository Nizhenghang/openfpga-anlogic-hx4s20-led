from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = DOCS / "基于OpenFPGA与安路HX4S20的自定义FPGA流水灯实验报告.docx"
DIAGRAM = DOCS / "custom_fpga_structure_for_report.png"


def font(size=22, bold=False):
    candidates = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for item in candidates:
        if Path(item).exists():
            return ImageFont.truetype(item, size=size)
    return ImageFont.load_default()


def draw_center(draw, box, text, fnt, fill=(15, 23, 42), spacing=6):
    x0, y0, x1, y1 = box
    lines = text.split("\n")
    line_heights = []
    widths = []
    for line in lines:
        b = draw.textbbox((0, 0), line, font=fnt)
        widths.append(b[2] - b[0])
        line_heights.append(b[3] - b[1])
    total_h = sum(line_heights) + spacing * (len(lines) - 1)
    y = y0 + ((y1 - y0) - total_h) / 2
    for line, w, h in zip(lines, widths, line_heights):
        x = x0 + ((x1 - x0) - w) / 2
        draw.text((x, y), line, font=fnt, fill=fill)
        y += h + spacing


def rounded_box(draw, box, fill, outline, width=3, radius=18):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, fill=(71, 85, 105), width=4):
    draw.line([start, end], fill=fill, width=width)
    x0, y0 = start
    x1, y1 = end
    if abs(x1 - x0) > abs(y1 - y0):
        direction = 1 if x1 > x0 else -1
        pts = [(x1, y1), (x1 - 14 * direction, y1 - 8), (x1 - 14 * direction, y1 + 8)]
    else:
        direction = 1 if y1 > y0 else -1
        pts = [(x1, y1), (x1 - 8, y1 - 14 * direction), (x1 + 8, y1 - 14 * direction)]
    draw.polygon(pts, fill=fill)


def make_diagram():
    img = Image.new("RGB", (1800, 1180), "#ffffff")
    draw = ImageDraw.Draw(img)
    title_font = font(42, True)
    h_font = font(28, True)
    body_font = font(23)
    small_font = font(20)

    draw.text((70, 48), "自定义 FPGA 在安路 HX4S20 上的层次结构", font=title_font, fill="#0f172a")
    draw.line((70, 110, 1730, 110), fill="#2563eb", width=4)

    board = (80, 160, 1720, 1080)
    wrapper = (180, 270, 1620, 1010)
    fabric = (315, 465, 1485, 920)
    rounded_box(draw, board, "#eef6ff", "#2563eb", width=4, radius=26)
    rounded_box(draw, wrapper, "#f8fafc", "#64748b", width=4, radius=24)
    rounded_box(draw, fabric, "#fff7ed", "#ea580c", width=4, radius=22)

    draw_center(draw, (120, 178, 1680, 235), "物理承载平台：Anlogic HX4S20 FPGA 开发板", h_font)
    draw_center(draw, (220, 292, 1580, 350), "板级封装：anlogic_openfpga_wrapper", h_font)
    draw_center(draw, (350, 487, 1450, 542), "OpenFPGA 生成的自定义 FPGA fabric：fpga_top", h_font)

    boxes = [
        ((250, 370, 560, 445), "配置控制\nRESET / LOAD / RUN", "#e0f2fe", "#0284c7"),
        ((620, 370, 930, 445), "内置配置数据\nfabric_bitstream.bit", "#f5f3ff", "#7c3aed"),
        ((990, 370, 1300, 445), "运行时钟分频\nLED 速度可见", "#ecfccb", "#65a30d"),
        ((1360, 370, 1560, 445), "LED 映射", "#fef9c3", "#ca8a04"),
        ((385, 585, 660, 700), "CCFF 配置链\n写入 LUT 与路由选择", "#f5f3ff", "#7c3aed"),
        ((735, 585, 1015, 700), "CLB / FLE\nLUT + FF", "#dcfce7", "#16a34a"),
        ((1090, 585, 1415, 700), "可配置互连\nSB / CBX / CBY / Channel", "#fef9c3", "#ca8a04"),
        ((560, 760, 1240, 855), "被配置后的用户电路：led_shift one-hot 流水灯", "#fee2e2", "#dc2626"),
    ]

    for b, text, fill, outline in boxes:
        rounded_box(draw, b, fill, outline, width=3, radius=16)
        draw_center(draw, b, text, body_font)

    arrow(draw, (560, 408), (620, 408))
    arrow(draw, (930, 408), (990, 408))
    arrow(draw, (1300, 408), (1360, 408))
    arrow(draw, (775, 445), (555, 585))
    arrow(draw, (775, 445), (875, 585))
    arrow(draw, (875, 700), (875, 760))
    arrow(draw, (1252, 700), (1120, 760))
    arrow(draw, (1240, 808), (1360, 408))

    pin_font = font(21)
    pin_box1 = (230, 955, 670, 1030)
    pin_box2 = (1130, 955, 1570, 1030)
    rounded_box(draw, pin_box1, "#ffffff", "#94a3b8", width=2, radius=14)
    rounded_box(draw, pin_box2, "#ffffff", "#94a3b8", width=2, radius=14)
    draw_center(draw, pin_box1, "输入：CLOCK / R7", pin_font)
    draw_center(draw, pin_box2, "输出：LED[3:0] / A4 A3 C10 B12", pin_font)
    arrow(draw, (450, 955), (450, 445))
    arrow(draw, (1360, 445), (1360, 955))

    draw.text(
        (90, 1110),
        "说明：HX4S20 负责承载整个 soft FPGA；流水灯逻辑通过 OpenFPGA bitstream 配置到 fabric 内部运行。",
        font=small_font,
        fill="#334155",
    )
    img.save(DIAGRAM)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, pct=5000):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), str(pct))


def style_run(run, size=None, bold=None, color=None, font_name="Microsoft YaHei"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text="", style=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        style_run(run, 11.5)
    return p


def add_body(doc, text):
    p = add_para(doc)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.18
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(text)
    style_run(run, 11.5)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        style_run(run, 16 if level == 1 else 13.5 if level == 2 else 12, True, (15, 23, 42))
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.9)
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(item)
        style_run(run, 11.5)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.9)
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(item)
        style_run(run, 11.5)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "EAF2F8")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    style_run(run, 10.5)
    if widths:
        for row in table.rows:
            for idx, w in enumerate(widths):
                row.cells[idx].width = Cm(w)
    doc.add_paragraph()
    return table


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)
    doc.add_paragraph()


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    style_run(run, 10.5, False, (71, 85, 105))


def build_docx():
    make_diagram()
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.3)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.6)
    sec.right_margin = Cm(2.4)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(11.5)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Microsoft YaHei"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    header = sec.header.paragraphs[0]
    header.text = "基于 OpenFPGA 与安路 HX4S20 的自定义 FPGA 实验报告"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        style_run(run, 9.5, False, (100, 116, 139))

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "OpenFPGA 自定义 FPGA 流水灯验证"
    for run in footer.runs:
        style_run(run, 9.5, False, (100, 116, 139))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("基于 OpenFPGA 与安路 HX4S20 的\n自定义 FPGA 流水灯实验报告")
    style_run(r, 22, True, (15, 23, 42))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("实验类型：FPGA 架构生成、软 FPGA 部署与板级验证")
    style_run(r, 12.5, False, (51, 65, 85))

    metadata = [
        ("项目名称", "OpenFPGA on Anlogic HX4S20 LED Demo"),
        ("硬件平台", "安路 HX4S20 FPGA 开发板"),
        ("软件工具", "WSL、OpenFPGA、Anlogic TD、Git/GitHub"),
        ("实验现象", "4 路 LED 按 one-hot 方式依次点亮，形成流水灯效果"),
        ("开源地址", "https://github.com/Nizhenghang/openfpga-anlogic-hx4s20-led"),
    ]
    table = doc.add_table(rows=len(metadata), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    for i, (k, v) in enumerate(metadata):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
        set_cell_shading(table.cell(i, 0), "EAF2F8")
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                for run in p.runs:
                    style_run(run, 10.5, bold=(cell == row.cells[0]))

    doc.add_page_break()

    add_heading(doc, "一、实验背景与意义")
    add_body(doc, "FPGA 的核心价值在于硬件结构可以在出厂后由用户重新配置。传统 FPGA 开发关注的是如何使用某一厂商已经设计好的 FPGA 芯片资源，而本实验进一步关注 FPGA 结构本身的生成和验证。通过 OpenFPGA，可以从体系结构描述出发生成一套可综合的 FPGA fabric，并在该 fabric 上映射用户电路。")
    add_body(doc, "本实验选择 LED 流水灯作为验证对象，是因为流水灯电路结构简单、实验现象直观，适合作为自定义 FPGA 部署流程的第一阶段验证。需要强调的是，本实验的重点不是流水灯逻辑本身，而是证明一条完整链路可以闭环：自定义 FPGA 架构生成、用户电路映射、配置 bitstream 产生、soft FPGA 综合到国产 FPGA、板级下载运行。")
    add_body(doc, "在国产 FPGA 平台上完成该流程具有实际意义。一方面可以验证 OpenFPGA 生成结果对不同厂商工具链的适配能力；另一方面也能帮助理解真实 FPGA 内部资源与 soft FPGA fabric 之间的层次关系，为后续扩展更复杂用户电路、修改 FPGA 架构参数和跨平台部署打下基础。")

    add_heading(doc, "二、实验目的")
    add_bullets(doc, [
        "理解 FPGA 中可配置逻辑块、查找表、触发器、布线网络和配置链的作用。",
        "掌握 OpenFPGA 生成自定义 FPGA fabric 的基本流程，并识别其关键输出文件。",
        "将 LED 流水灯用户电路映射到自定义 FPGA fabric 内部，而不是直接综合到物理 FPGA。",
        "在安路 TD 工具中完成 Verilog 导入、顶层设置、管脚约束、综合、布局布线和下载验证。",
        "通过 one-hot LED 流水灯现象验证自定义 FPGA 配置链和用户电路运行正确。",
    ])

    add_heading(doc, "三、实验环境")
    add_heading(doc, "3.1 硬件环境", 2)
    add_body(doc, "硬件平台为安路 HX4S20 FPGA 开发板。该开发板提供板载时钟输入和 4 个可观察 LED 输出。本实验将 OpenFPGA 生成的自定义 FPGA fabric 作为一段普通可综合 Verilog 逻辑部署到 HX4S20 芯片内部，HX4S20 因此扮演承载 soft FPGA 的物理平台。")
    add_table(doc, ["信号", "管脚", "功能说明"], [
        ("CLOCK", "R7", "板载时钟输入，用作 wrapper 配置状态机与运行时钟来源"),
        ("LED[0]", "A4", "流水灯输出 0"),
        ("LED[1]", "A3", "流水灯输出 1"),
        ("LED[2]", "C10", "流水灯输出 2"),
        ("LED[3]", "B12", "流水灯输出 3"),
    ], [3.2, 3.2, 9.0])

    add_heading(doc, "3.2 软件环境", 2)
    add_table(doc, ["工具", "在实验中的作用"], [
        ("WSL", "用于运行 OpenFPGA flow，完成用户电路映射和 fabric 相关文件生成"),
        ("OpenFPGA", "生成自定义 FPGA fabric、配置链、bitstream 和仿真/验证相关 netlist"),
        ("Anlogic TD", "完成面向 HX4S20 的综合、布局布线、bitstream 生成和板卡下载"),
        ("Git / GitHub", "整理实验工程、保存关键文件并发布开源仓库"),
    ], [4.0, 11.4])

    add_heading(doc, "四、实验总体方案")
    add_body(doc, "实验采用分层实现方案。最底层是安路 HX4S20 物理 FPGA，它提供真实 LUT、触发器、布线、IO 和下载能力。中间层是 OpenFPGA 生成的自定义 FPGA fabric，该 fabric 本质上是一套由 Verilog 描述的可配置硬件结构。最上层是用户流水灯电路，该电路通过 OpenFPGA 的映射流程变成 fabric bitstream，最终写入 soft FPGA 内部运行。")
    add_body(doc, "因此，TD 工具综合的并不是单纯的流水灯 RTL，而是一个已经包含 CLB、互连网络和配置链的可配置平台。流水灯只是这个平台被配置后的功能表现。通过这种方式，可以直观看到自定义 FPGA 与传统 FPGA 应用之间的区别。")
    doc.add_picture(str(DIAGRAM), width=Inches(6.3))
    add_caption(doc, "图 1 自定义 FPGA 在 HX4S20 上的部署层次")

    add_heading(doc, "五、实验原理")
    add_heading(doc, "5.1 普通流水灯与本实验流水灯的差异", 2)
    add_body(doc, "普通 FPGA 流水灯实验通常直接编写计数器和移位寄存器，厂商工具会把这些 RTL 逻辑直接映射到物理 FPGA 的原生 LUT 和触发器中。该方式流程短、资源占用低，但无法体现 FPGA fabric 本身的结构。")
    add_body(doc, "本实验采用两级映射。第一步，OpenFPGA 将用户流水灯电路映射到自定义 FPGA fabric 的逻辑块和互连资源中，并生成 fabric 配置 bitstream。第二步，TD 将整个 soft FPGA fabric 综合到 HX4S20 芯片中。最终 LED 的变化来自 soft FPGA 内部被配置后的用户逻辑，而不是 TD 直接综合出来的普通流水灯模块。")

    add_heading(doc, "5.2 自定义 FPGA fabric 的组成", 2)
    add_body(doc, "本实验使用的自定义 FPGA fabric 主要由可配置逻辑资源、可配置互连资源、IO 结构和配置链组成。可配置逻辑资源负责实现用户电路中的组合逻辑和时序逻辑；互连资源负责在不同逻辑块之间建立连接；IO 结构负责将 fabric 内部信号连接到 wrapper；配置链负责将 bitstream 写入 fabric 内部配置存储单元。")
    add_bullets(doc, [
        "CLB：可配置逻辑块，是 soft FPGA 的核心计算资源。",
        "FLE / LUT / FF：用于实现布尔逻辑函数和寄存器。",
        "Switch Block 与 Connection Block：控制路由通道和逻辑块端口之间的连接关系。",
        "Routing Channel：提供横向和纵向布线轨道。",
        "CCFF 配置链：保存 LUT 内容、mux 选择和互连配置信息。",
    ])

    add_heading(doc, "5.3 wrapper 的作用", 2)
    add_body(doc, "`anlogic_openfpga_wrapper` 是 TD 工程中的顶层模块。它不是 OpenFPGA fabric 本身，而是物理开发板和 soft FPGA 之间的适配层。该模块负责接收板载时钟、产生较慢的运行时钟、控制配置状态机、将内置 bitstream 写入 CCFF 配置链，并把 fabric 输出映射到板载 LED。")
    add_body(doc, "顶层模块必须设置为 `anlogic_openfpga_wrapper`，原因是 TD 需要从该模块开始展开整个设计。若误将 `fpga_top` 或其他内部模块设为顶层，则板级管脚、配置加载、时钟分频和 LED 映射都无法正确连接，下载后也就无法出现预期实验现象。")

    add_heading(doc, "六、实验工程文件说明")
    add_body(doc, "实验工程已经整理为适合开源和复现实验的目录结构。TD 工程只需要导入 `td/` 目录下的 Verilog 文件和约束文件，其他目录用于保存 OpenFPGA 关键产物、说明文档和原生流水灯对照测试。")
    add_code_block(doc, """openfpga-anlogic-hx4s20-led/
├── td/
│   ├── anlogic_openfpga_onehot_led.v
│   └── anlogic_openfpga_wrapper.adc
├── openfpga_artifacts/
│   ├── fpga_top.v
│   ├── fabric_bitstream.bit
│   └── fabric_independent_bitstream.xml
├── native_led_test/
│   ├── hx4s20_led_chase.v
│   └── hx4s20_led_chase.adc
└── docs/""")
    add_table(doc, ["文件", "说明"], [
        ("td/anlogic_openfpga_onehot_led.v", "TD 中实际导入的最终 Verilog 文件，已经包含 wrapper 和 OpenFPGA fabric 相关逻辑"),
        ("td/anlogic_openfpga_wrapper.adc", "HX4S20 板级管脚约束，指定时钟和 LED 管脚"),
        ("openfpga_artifacts/fpga_top.v", "OpenFPGA 生成的自定义 FPGA fabric 顶层，用于说明 fabric 结构来源"),
        ("openfpga_artifacts/fabric_bitstream.bit", "配置 soft FPGA 的 bitstream"),
        ("openfpga_artifacts/fabric_independent_bitstream.xml", "OpenFPGA 生成的 bitstream 描述文件"),
        ("native_led_test/", "普通原生流水灯对照测试，用于区分直接 FPGA 实现和 soft FPGA 实现"),
    ], [5.2, 10.2])

    add_heading(doc, "七、实验步骤")
    add_heading(doc, "7.1 用户电路设计", 2)
    add_body(doc, "用户电路设计为 4 路 LED 流水灯。为了使实验现象明确，最终显示方式调整为 one-hot，即每个时刻主要只有一个 LED 点亮。该设计能直观反映计数、状态切换和输出译码是否正常。")

    add_heading(doc, "7.2 OpenFPGA 生成与映射", 2)
    add_body(doc, "在 WSL 环境中运行 OpenFPGA flow，将 LED 流水灯用户电路映射到指定的自定义 FPGA 架构。OpenFPGA 完成逻辑综合、技术映射、布局、布线和 bitstream 生成。该阶段产生的 `fpga_top.v` 表示 soft FPGA fabric 的硬件结构，`fabric_bitstream.bit` 表示把流水灯用户电路配置进 fabric 所需的数据。")

    add_heading(doc, "7.3 TD 兼容化处理", 2)
    add_body(doc, "OpenFPGA 原始输出文件并不能完全按原样导入 TD。一方面，多文件 include 结构可能导致 TD 路径解析不稳定；另一方面，部分 pass-gate mux、三态或高阻相关写法对 TD 综合不够友好。因此实验中将 OpenFPGA 相关 netlist 合并为单个 TD 专用 Verilog 文件，并对不兼容结构进行改写。")
    add_body(doc, "该步骤是实验成功的关键。如果把 OpenFPGA 生成目录中的测试平台、旧 wrapper 或中间文件全部加入 TD，很容易出现顶层混乱、重复定义、综合错误或下载后无现象。因此最终规定 TD 只导入一个 Verilog 源文件和一个 ADC 约束文件。")

    add_heading(doc, "7.4 TD 工程建立与顶层设置", 2)
    add_numbered(doc, [
        "在 TD 中新建工程，选择与开发板一致的 HX4S20 器件型号。",
        "添加 Verilog 源文件 `td/anlogic_openfpga_onehot_led.v`。",
        "设置顶层模块为 `anlogic_openfpga_wrapper`。",
        "添加约束文件 `td/anlogic_openfpga_wrapper.adc`。",
        "依次运行综合、布局布线、bitstream 生成和下载流程。",
    ])

    add_heading(doc, "7.5 管脚约束", 2)
    add_body(doc, "ADC 约束文件用于把顶层端口绑定到开发板实际管脚。如果没有导入约束文件，即使逻辑综合成功，时钟和 LED 也可能没有连接到正确物理管脚，实验现象就无法正常观察。")
    add_code_block(doc, """set_pin_assignment { CLOCK } { LOCATION = R7; IOSTANDARD = LVCMOS33; PULLTYPE = PULLUP; }
set_pin_assignment { LED[0] } { LOCATION = A4; IOSTANDARD = LVCMOS33; DRIVESTRENGTH = 8; PULLTYPE = NONE; }
set_pin_assignment { LED[1] } { LOCATION = A3; IOSTANDARD = LVCMOS33; DRIVESTRENGTH = 8; PULLTYPE = NONE; }
set_pin_assignment { LED[2] } { LOCATION = C10; IOSTANDARD = LVCMOS33; DRIVESTRENGTH = 8; PULLTYPE = NONE; }
set_pin_assignment { LED[3] } { LOCATION = B12; IOSTANDARD = LVCMOS33; DRIVESTRENGTH = 8; PULLTYPE = NONE; }""")

    add_heading(doc, "八、实验结果与资源占用")
    add_heading(doc, "8.1 实验现象", 2)
    add_body(doc, "工程下载到 HX4S20 开发板后，4 个板载 LED 按固定顺序依次闪烁。调整显示逻辑后，实验现象满足 one-hot 流水灯要求，即任意时刻主要只有一个 LED 被点亮，然后依次切换到下一个 LED。")
    add_body(doc, "该现象说明物理 FPGA 中的 wrapper 已经完成 soft FPGA 配置加载，OpenFPGA 生成的配置 bitstream 已写入 CCFF 配置链，用户流水灯逻辑已经在自定义 FPGA fabric 内部运行并通过 IO 映射输出到开发板 LED。")

    add_heading(doc, "8.2 资源占用", 2)
    add_table(doc, ["资源类型", "数量"], [
        ("LUT", "1741"),
        ("Sequential cells", "2043"),
        ("Pads", "5"),
        ("BRAM", "0"),
        ("DSP", "0"),
        ("Packed mslices", "631"),
        ("Packed lslices", "632"),
        ("面积利用率", "约 16%"),
    ], [6.0, 5.0])
    add_body(doc, "从资源占用可以看出，本实验消耗的逻辑资源明显高于普通 LED 流水灯。这是合理的，因为这里综合到 HX4S20 中的是整个 soft FPGA fabric，而不是单独的流水灯电路。资源主要用于实现 CLB 阵列、可配置互连、配置链、mux 结构、时钟分频和 wrapper 控制逻辑。")

    add_heading(doc, "九、问题分析与解决")
    add_heading(doc, "9.1 下载后没有出现流水灯", 2)
    add_body(doc, "最初下载到开发板后没有出现预期流水灯现象。排查时首先确认了 TD 工程的顶层模块设置、Verilog 文件导入情况和 ADC 约束是否生效。由于 OpenFPGA 生成文件较多，如果导入了错误版本或把测试平台加入综合工程，TD 可能并不会按预期的板级顶层进行综合。最终通过明确只导入 TD 专用单文件，并将顶层设为 `anlogic_openfpga_wrapper`，解决了顶层和文件选择问题。")

    add_heading(doc, "9.2 LED 不是单灯依次点亮", 2)
    add_body(doc, "实验中曾观察到 4 个 LED 有规律闪烁，但并不是每次只亮一个灯。这说明系统内部时钟、配置和输出路径已经部分工作，但 LED 显示译码方式不符合最终需求。后续将 LED 映射调整为 one-hot 方式，使实验现象更直观地表现为单灯依次移动。")

    add_heading(doc, "9.3 TD 与 OpenFPGA 输出兼容性", 2)
    add_body(doc, "OpenFPGA 生成的 Verilog 面向通用仿真和综合流程，不能默认认为所有 FPGA 厂商工具都能无修改接受。TD 对部分结构的支持与其他工具存在差异，因此需要进行兼容化处理。实验中主要处理了三态/高阻、pass-gate 风格 mux、测试平台代码和多文件 include 等问题。这说明在跨工具链部署 soft FPGA 时，工程整理和兼容性适配与架构生成同样重要。")

    add_heading(doc, "十、可移植性分析")
    add_body(doc, "从理论上讲，本实验中的自定义 FPGA fabric 是普通 Verilog 描述，因此可以迁移到其他厂商 FPGA 上。但是迁移不是简单复制工程。不同厂商工具在约束格式、IO 标准、时钟资源、综合支持语法、下载流程和时序约束方面均存在差异。")
    add_body(doc, "如果部署到 Xilinx、Intel、Lattice 或其他平台，需要重新编写对应约束文件，例如 Xilinx 使用 XDC，Intel 使用 QSF/SDC；同时还需要检查 wrapper 中时钟、复位、IO buffer 和配置控制逻辑是否符合目标工具要求。若目标工具不支持当前 Verilog 中的某些结构，还需要再次进行兼容化改写。")
    add_body(doc, "因此，本项目具备跨平台迁移的基础，但不能保证零修改直接部署。比较合理的迁移策略是保持 OpenFPGA fabric 和用户 bitstream 的核心逻辑不变，针对目标开发板重新实现板级 wrapper、管脚约束和工具兼容层。")

    add_heading(doc, "十一、实验结论")
    add_body(doc, "本实验完成了从 OpenFPGA 自定义 FPGA 结构生成，到安路 HX4S20 平台综合、布局布线和板级下载验证的完整流程。实验最终实现了 4 路 LED one-hot 流水灯，证明 OpenFPGA 生成的 soft FPGA fabric 可以部署到国产安路 FPGA 上运行。")
    add_body(doc, "通过本实验可以明确区分两类设计：普通 FPGA 应用设计关注的是在现成 FPGA 上实现某个功能；本实验关注的是先构造一个可配置 FPGA fabric，再把用户电路配置到这个 fabric 内部运行。前者是使用 FPGA，后者更接近设计并验证一个小型 FPGA。")
    add_body(doc, "实验结果表明，OpenFPGA 的生成流程、TD 的综合实现流程以及安路开发板的下载验证流程可以形成闭环。后续可以在该基础上进一步修改自定义 FPGA 架构规模，替换更复杂的用户电路，或尝试迁移到其他 FPGA 厂商平台。")

    add_heading(doc, "十二、项目开源说明")
    add_body(doc, "本实验工程已经整理并上传到 GitHub。仓库中保留了 TD 可导入文件、OpenFPGA 关键生成物、资源占用说明、结构示意图和实验流程文档。为了避免版权风险，仓库没有公开上传厂商数据手册 PDF。")
    add_code_block(doc, "https://github.com/Nizhenghang/openfpga-anlogic-hx4s20-led")

    doc.save(OUT)


if __name__ == "__main__":
    build_docx()
    print(OUT)
